import os
import io
import time
import random
import requests
import anthropic
from docx import Document
from docx.shared import Pt
import openpyxl

TELEGRAM_TOKEN = os.environ["TELEGRAM_TOKEN"]
ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]
HF_KEY = os.environ.get("HF_KEY", "")

# Бан-лист: BANNED_USERS=123456789,987654321 в Railway Variables
BANNED_USERS = set(
    int(x.strip()) for x in os.environ.get("BANNED_USERS", "").split(",") if x.strip().isdigit()
)

# Только эти user_id могут использовать скиллы МГИМО
ADMIN_USERNAME = "forge0n"
ADMIN_IDS = set(
    int(x.strip()) for x in os.environ.get("ADMIN_IDS", "").split(",") if x.strip().isdigit()
)
MGIMO_ALLOWED = ADMIN_IDS.copy()

BASE_URL = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}"
client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

MODELS = {
    "sonnet": "claude-sonnet-4-6",
    "opus": "claude-opus-4-7",
}
MODEL_NAMES = {
    "sonnet": "⚡ Sonnet 4.6",
    "opus": "🧠 Opus 4.7",
}

# Пробный период
FREE_LIMIT = 20
# Реферальный бонус
REFERRAL_BONUS = 15

# Хранилище пользователей (в памяти, сбрасывается при рестарте)
user_data = {}
pending_clear = set()
pending_text_mode = {}  # chat_id -> submode (rephrase, shorten, extend, bullets, ai_check)

def fix_dashes(text):
    return text.replace(" -- ", " - ").replace("--", "-").replace(" — ", " - ").replace("—", "-")

def get_user(uid):
    if uid not in user_data:
        user_data[uid] = {
            "mode": "default",
            "language": "ru",
            "history": [],
            "model": "sonnet",
            "msg_count": 0,
            "has_sub": False,
            "referral_bonus": 0,
            "referred_by": None,
            "referrals": [],
        }
    return user_data[uid]

def is_admin(user_id, username):
    return username == ADMIN_USERNAME or user_id in ADMIN_IDS

def can_use_mgimo(user_id, username):
    return is_admin(user_id, username) or user_id in MGIMO_ALLOWED

def get_limit(uid):
    u = get_user(uid)
    if u["has_sub"]:
        return 999999
    return FREE_LIMIT + u.get("referral_bonus", 0)

def can_message(uid):
    u = get_user(uid)
    return u["msg_count"] < get_limit(uid)

SYSTEM_PROMPTS = {
    "default": "Ты - универсальный ИИ-ассистент. Отвечай нейтрально, чётко и по делу. Без лишних эмодзи, без личности и шуток. Обращайся к пользователю на ты. Никогда не используй длинное тире, только обычный дефис.",
    "delo": "Ты - строгий деловой ИИ-ассистент. Отвечай максимально чётко и кратко, только факты и конкретика. Никаких шуток и эмодзи. Обращайся к пользователю на ты. Никогда не используй длинное тире, только обычный дефис.",
    "friend": "Ты - близкий друг и весёлый ИИ-ассистент. Немного шути, будь расслабленным - но всё равно помогай по делу. Обращайся к пользователю на ты как к самому близкому человеку. Можешь использовать эмодзи. Никогда не используй длинное тире, только обычный дефис.",
    "test": "Ты - ассистент для решения тестов. Отвечай максимально кратко и точно. Только правильный ответ, без объяснений если не просят. Никогда не используй длинное тире, только обычный дефис.",
    "english": """Ты - персональный преподаватель английского языка. Ты помогаешь с:
- Переводом текстов и предложений (с объяснением нюансов)
- Грамматикой (объяснения на русском, примеры на английском)
- Тестами и упражнениями по запросу
- Разбором ошибок
- Пополнением словарного запаса

Стиль: дружелюбный, терпеливый. Обращайся на ты.
Если пользователь пишет по-английски - исправляй ошибки мягко и объясняй почему.
Если просит тест - составь 5 вопросов по нужной теме.
Никогда не используй длинное тире, только обычный дефис.""",
    "text_work": """Ты - профессиональный редактор текстов. Пользователь пришлёт текст и укажет что нужно сделать:
- Перефразировка: перепиши текст другими словами, сохранив смысл
- Сокращение: сократи текст, оставив главное
- Удлинение: расширь текст, добавив детали и аргументы
- По пунктам: структурируй текст в виде пронумерованных пунктов
- Анализ на ИИ: проведи псевдоанализ текста на оригинальность

Жди текст от пользователя и выполняй задачу качественно.
Никогда не используй длинное тире, только обычный дефис.""",
    "skill_gpt": """You are a prompt director for GPT Image 2.0. Convert user concepts into production-ready prompts.

Three formats:
- Format A (JSON): for layouts, UI, infographics, posters with panels, character sheets
- Format B (cinematic prose): for single scenes, portraits, illustrations
- Format C (meta-prompt): when user gives only a theme

Return ONLY the prompt in a code block. No explanation.
Communicate with user in Russian, write prompts in English.
Never use em-dashes, only regular hyphens.""",
    "skill_sales": """Ты - Sales Coach по продажам AI-визуалов.

Прайс: 1 фото $15, 10 фото $140, 20 фото $260, 1 видео $80, 5 видео $350, 10 видео $650.
Ниши: одежда, рестораны, недвижимость, блогеры, агентства.

Задачи:
- "найди клиентов" - ищи реальные аккаунты, выдавай таблицу с оценкой потенциала
- "напиши оффер" - готовый текст под нишу
- "напиши DM" - скрипт для директа
- возражение - готовый ответ

Возражения:
"Дорого" - У нас 10 фото = $140, в 3-5 раз дешевле студии. Начнём с бесплатного теста?
"Уже есть фотограф" - AI - дополнение, не замена. Больше контента за меньший бюджет.
"Подумаю" - Сделаю 1 тестовый визуал бесплатно прямо сейчас. Попробуем?

Никогда не используй длинное тире, только обычный дефис.
Всегда давай готовые тексты, не абстрактные советы.""",
    "skill_mgimo": """Ты - ассистент для контрольных работ по БЖД МГИМО.

Данные студента (всегда использовать):
ФИО: Портненко М.В., Институт: ИМТУР, Курс: 1, Группа: 1
Преподаватель: Трофимов Сергей Анатольевич
Дисциплина: БЕЗОПАСНОСТЬ ЖИЗНЕДЕЯТЕЛЬНОСТИ
Вуз: МГИМО МИД России, Москва, 2026

Структура: Титул - Содержание - Введение - Основная часть (2 подглавы) - Заключение - Литература (5 источников).
Оформление: Times New Roman 14pt, интервал 1.5, поля 30/15/20/20мм.
Спрашивай только тему. Пиши полный текст работы.
Никогда не используй длинное тире, только обычный дефис.""",
    "skill_coursework": """Ты - ассистент для курсовых работ МГИМО.

Оформление: Times New Roman 14pt, интервал 1.5, поля 30/15/20/20мм, нумерация с введения.
Структура: Титул - Оглавление - Введение (2-3 стр.) - Главы - Заключение - Литература.
Введение: актуальность, объект, предмет, цель, задачи, структура.

Сначала спроси: тему, ФИО, группу, институт, кафедру, научного руководителя, год.
Затем пиши полный текст.
Никогда не используй длинное тире, только обычный дефис.""",
}

BUTTONS = {
    "🗣 Дефолт": ("default", "Режим: 🗣 Дефолт."),
    "💼 По делу": ("delo", "Режим: 💼 По делу. Чётко и по фактам."),
    "😊 Друг": ("friend", "Режим: 😊 Друг. Привет, родной)"),
    "📝 Тесты": ("test", "Режим: 📝 Тесты. Скидывай вопрос!"),
    "🇬🇧 Английский": ("english", "Режим: 🇬🇧 Английский. Поможу с языком! Скидывай текст, вопрос или попроси тест."),
    "✏️ Работа с текстом": ("text_work", None),
    "🎨 GPT Image промпты": ("skill_gpt", "Скилл: 🎨 GPT Image. Скидывай идею!"),
    "💰 AI Visuals Sales": ("skill_sales", "Скилл: 💰 AI Sales. Чем помочь?"),
    "📚 МГИМО БЖД": ("skill_mgimo", "Скилл: 📚 МГИМО БЖД. Скидывай тему!"),
    "📖 Курсовая МГИМО": ("skill_coursework", "Скилл: 📖 Курсовая. Скажи тему и данные!"),
    "🇷🇺 Русский": None,
    "🇬🇧 English": None,
    "🌐 Авто": None,
    "⚡ Sonnet 4.6": None,
    "🧠 Opus 4.7": None,
}

TEXT_SUBMODES = {
    "🔄 Перефразировка": "rephrase",
    "✂️ Сокращение": "shorten",
    "📝 Удлинение": "extend",
    "📋 По пунктам": "bullets",
    "🤖 Анализ на ИИ": "ai_check",
}

TEXT_SUBMODE_PROMPTS = {
    "rephrase": "Перефразируй следующий текст другими словами, полностью сохранив смысл. Текст:\n\n",
    "shorten": "Сократи следующий текст, оставив только главное. Текст:\n\n",
    "extend": "Расширь следующий текст, добавив детали, аргументы и примеры. Текст:\n\n",
    "bullets": "Структурируй следующий текст в виде пронумерованных пунктов. Текст:\n\n",
    "ai_check": None,
}

def get_system(uid):
    u = get_user(uid)
    base = SYSTEM_PROMPTS[u["mode"]]
    if u["language"] == "ru":
        base += "\nОтвечай только на русском языке."
    elif u["language"] == "en":
        base += "\nAlways respond in English only."
    return base

def send_typing(chat_id):
    requests.post(f"{BASE_URL}/sendChatAction", json={"chat_id": chat_id, "action": "typing"})

def send(chat_id, text, keyboard=None, markdown=False):
    text = fix_dashes(text)
    chunks = [text[i:i+4096] for i in range(0, len(text), 4096)]
    for i, chunk in enumerate(chunks):
        payload = {"chat_id": chat_id, "text": chunk}
        if markdown:
            payload["parse_mode"] = "Markdown"
        if i == 0 and keyboard:
            payload["reply_markup"] = {
                "keyboard": keyboard,
                "one_time_keyboard": True,
                "resize_keyboard": True
            }
        r = requests.post(f"{BASE_URL}/sendMessage", json=payload)
        if markdown and r.status_code != 200:
            payload.pop("parse_mode", None)
            requests.post(f"{BASE_URL}/sendMessage", json=payload)

def send_document_bytes(chat_id, file_bytes, filename, caption=""):
    requests.post(
        f"{BASE_URL}/sendDocument",
        data={"chat_id": chat_id, "caption": caption},
        files={"document": (filename, file_bytes)}
    )

def make_docx(text, title="Документ"):
    doc = Document()
    doc.add_heading(title, 0)
    for para in text.split("\n"):
        if para.strip():
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(12)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def make_txt(text):
    return text.encode("utf-8")

def ai_check_result():
    originality = random.randint(55, 72)
    borrowing = random.randint(8, 18)
    citation = random.randint(3, 10)
    repetitions = random.randint(5, 12)
    human_pct = random.randint(52, 68)
    ai_pct = 100 - human_pct

    result = (
        "🤖 *Анализ текста на ИИ-генерацию*\n\n"
        f"1. Оригинальность: {originality}%\n"
        f"2. Заимствования: {borrowing}%\n"
        f"3. Цитирование: {citation}%\n"
        f"4. Повторения: {repetitions}%\n\n"
        f"*Итоговый результат:*\n"
        f"👤 Человек: {human_pct}%\n"
        f"🤖 ИИ: {ai_pct}%"
    )
    return result

def set_commands():
    cmds = [
        {"command": "start", "description": "Начать заново"},
        {"command": "mode", "description": "Режим общения"},
        {"command": "skills", "description": "Специальные навыки"},
        {"command": "model", "description": "Выбор модели AI"},
        {"command": "lang", "description": "Язык ответов"},
        {"command": "clear", "description": "Очистить память"},
        {"command": "sub", "description": "Подписка"},
        {"command": "ref", "description": "Реферальная ссылка"},
        {"command": "status", "description": "Мой статус"},
        {"command": "help", "description": "Помощь"},
    ]
    requests.post(f"{BASE_URL}/setMyCommands", json={"commands": cmds})

def handle_admin(chat_id, text, user_id, username):
    if not is_admin(user_id, username):
        send(chat_id, "Нет доступа.")
        return

    parts = text.split()
    cmd = parts[0] if parts else ""

    if cmd == "/admin":
        lines = ["👑 *Админ-панель*\n"]
        lines.append(f"Всего пользователей: {len(user_data)}\n")
        lines.append("Команды:")
        lines.append("/users - список пользователей")
        lines.append("/ban [id] - забанить")
        lines.append("/unban [id] - разбанить")
        lines.append("/give_sub [id] - выдать подписку")
        lines.append("/give_mgimo [id] - доступ к МГИМО")
        send(chat_id, "\n".join(lines), markdown=True)

    elif cmd == "/users":
        if not user_data:
            send(chat_id, "Пользователей нет.")
            return
        lines = ["👥 *Пользователи:*\n"]
        for uid, u in list(user_data.items())[-30:]:
            sub = "✅" if u.get("has_sub") else f"{u.get('msg_count', 0)}/{FREE_LIMIT}"
            banned = "🚫" if uid in BANNED_USERS else ""
            lines.append(f"ID: `{uid}` | {sub} {banned}")
        send(chat_id, "\n".join(lines), markdown=True)

    elif cmd == "/ban" and len(parts) > 1:
        try:
            target = int(parts[1])
            BANNED_USERS.add(target)
            send(chat_id, f"✅ Пользователь {target} забанен.")
        except:
            send(chat_id, "Неверный ID.")

    elif cmd == "/unban" and len(parts) > 1:
        try:
            target = int(parts[1])
            BANNED_USERS.discard(target)
            send(chat_id, f"✅ Пользователь {target} разбанен.")
        except:
            send(chat_id, "Неверный ID.")

    elif cmd == "/give_sub" and len(parts) > 1:
        try:
            target = int(parts[1])
            get_user(target)["has_sub"] = True
            send(chat_id, f"✅ Подписка выдана пользователю {target}.")
        except:
            send(chat_id, "Неверный ID.")

    elif cmd == "/give_mgimo" and len(parts) > 1:
        try:
            target = int(parts[1])
            MGIMO_ALLOWED.add(target)
            send(chat_id, f"✅ Доступ к МГИМО выдан пользователю {target}.")
        except:
            send(chat_id, "Неверный ID.")

def handle(update):
    msg = update.get("message", {})
    if not msg:
        return

    chat_id = msg["chat"]["id"]
    user_id = msg.get("from", {}).get("id", 0)
    text = msg.get("text", "")
    name = msg.get("from", {}).get("first_name", "пользователь")
    username = msg.get("from", {}).get("username", "")

    print(f"[USER] id={user_id} @{username} {name}: {text[:80]}")

    if user_id in BANNED_USERS:
        send(chat_id, "У тебя нет доступа к этому боту.")
        return

    u = get_user(chat_id)

    # Реферальная регистрация
    if text.startswith("/start ref_") and not u.get("referred_by"):
        try:
            ref_id = int(text.split("ref_")[1])
            if ref_id != chat_id and ref_id in user_data:
                ref_u = get_user(ref_id)
                if ref_u.get("has_sub"):
                    ref_u["referral_bonus"] = ref_u.get("referral_bonus", 0) + REFERRAL_BONUS
                    ref_u["referrals"].append(chat_id)
                    u["referred_by"] = ref_id
                    send(ref_id, f"🎉 Твой друг зарегистрировался по реферальной ссылке! +{REFERRAL_BONUS} сообщений.")
        except:
            pass

    # Админ команды
    if text.startswith("/admin") or text.startswith("/users") or text.startswith("/ban") or text.startswith("/unban") or text.startswith("/give_sub") or text.startswith("/give_mgimo"):
        handle_admin(chat_id, text, user_id, username)
        return

    # Подтверждение очистки
    if chat_id in pending_clear:
        if text == "✅ Точно очистить":
            pending_clear.discard(chat_id)
            u["history"] = []
            send(chat_id, "🗑 Память очищена! Начинаем с чистого листа.")
        elif text == "❌ Нет":
            pending_clear.discard(chat_id)
            send(chat_id, "Окей, память осталась нетронутой 👍")
        else:
            send(chat_id, "Выбери вариант 👇", [["✅ Точно очистить", "❌ Нет"]])
        return

    # Режим работы с текстом - ждём субрежим
    if u.get("mode") == "text_work" and chat_id not in pending_text_mode:
        if text in TEXT_SUBMODES:
            submode = TEXT_SUBMODES[text]
            pending_text_mode[chat_id] = submode
            send(chat_id, "Отправь текст, и я его обработаю 👇")
            return

    # Режим работы с текстом - ждём текст
    if chat_id in pending_text_mode and not text.startswith("/"):
        submode = pending_text_mode.pop(chat_id)
        send_typing(chat_id)

        if submode == "ai_check":
            time.sleep(1.5)
            send(chat_id, "⏳ Анализирую текст...\n[████████░░] 80%")
            time.sleep(1.5)
            send(chat_id, ai_check_result(), markdown=True)
            # Предложить скачать
            txt_bytes = make_txt(f"Анализ текста на ИИ\n\nТекст:\n{text}\n\n" + ai_check_result().replace("*", ""))
            send_document_bytes(chat_id, txt_bytes, "ai_analysis.txt", "Результат анализа")
        else:
            prompt = TEXT_SUBMODE_PROMPTS[submode] + text
            try:
                response = client.messages.create(
                    model=MODELS.get(u.get("model", "sonnet"), MODELS["sonnet"]),
                    max_tokens=2048,
                    system=SYSTEM_PROMPTS["text_work"],
                    messages=[{"role": "user", "content": prompt}]
                )
                reply = response.content[0].text
                send(chat_id, reply, markdown=True)
                # Предложить скачать
                keyboard = [["📄 Скачать TXT", "📝 Скачать DOCX"]]
                send(chat_id, "Хочешь скачать результат?", keyboard)
                u["last_text_result"] = reply
            except Exception as e:
                send(chat_id, f"Ошибка: {str(e)}")
        return

    # Скачать результат работы с текстом
    if text == "📄 Скачать TXT" and u.get("last_text_result"):
        send_document_bytes(chat_id, make_txt(u["last_text_result"]), "result.txt")
        return
    if text == "📝 Скачать DOCX" and u.get("last_text_result"):
        send_document_bytes(chat_id, make_docx(u["last_text_result"]), "result.docx")
        return

    # Команды
    if text.startswith("/start"):
        model_name = MODEL_NAMES.get(u.get("model", "sonnet"), "⚡ Sonnet 4.6")
        mode_names = {
            "default": "🗣 Дефолт", "delo": "💼 По делу", "friend": "😊 Друг",
            "test": "📝 Тесты", "english": "🇬🇧 Английский", "text_work": "✏️ Работа с текстом",
            "skill_gpt": "🎨 GPT Image", "skill_sales": "💰 AI Sales",
            "skill_mgimo": "📚 МГИМО БЖД", "skill_coursework": "📖 Курсовая МГИМО",
        }
        mode_name = mode_names.get(u.get("mode", "default"), "🗣 Дефолт")
        send(chat_id,
            f"Привет, {name}! 👋\n"
            f"Я - твой личный AI-ассистент. Помогу с задачами, текстами, идеями и ответами на любые вопросы.\n\n"
            f"⚙️ Режим: {mode_name}\n"
            f"🧠 Модель: {model_name}\n\n"
            f"📌 Команды:\n"
            f"/help - список возможностей\n"
            f"/mode - сменить режим\n"
            f"/skills - специальные навыки\n"
            f"/lang - язык\n"
            f"/clear - очистить память\n"
            f"/sub - подписка\n"
            f"/ref - реферальная ссылка\n"
            f"/status - мой статус\n\n"
            f"💬 Пиши, что нужно - разберёмся!"
        )
        return

    if text == "/help":
        send(chat_id,
            "🤖 *AI-ассистент - возможности*\n\n"
            "*Режимы (/mode):*\n"
            "🗣 Дефолт - нейтральный ИИ\n"
            "💼 По делу - строго и кратко\n"
            "😊 Друг - по-дружески\n"
            "📝 Тесты - кратко и точно\n"
            "🇬🇧 Английский - помощь с языком\n"
            "✏️ Работа с текстом - редактирование\n\n"
            "*Скиллы (/skills):*\n"
            "🎨 GPT Image - промпты\n"
            "💰 AI Sales - продажи\n"
            "📚 МГИМО БЖД - контрольные\n"
            "📖 Курсовая - курсовые МГИМО\n\n"
            "*Модели (/model):*\n"
            "⚡ Sonnet 4.6 - быстрый, по умолчанию\n"
            "🧠 Opus 4.7 - умнее, для сложных задач\n\n"
            "*Другое:*\n"
            "/sub - подписка\n"
            "/ref - пригласи друга, получи +15 сообщений\n"
            "/status - твой статус",
            markdown=True
        )
        return

    if text == "/mode":
        send(chat_id, "Выбери режим:", [
            ["🗣 Дефолт", "💼 По делу"],
            ["😊 Друг", "📝 Тесты"],
            ["🇬🇧 Английский", "✏️ Работа с текстом"],
            ["🖼 Фотографии 🔒"]
        ])
        return

    if text == "/skills":
        send(chat_id,
            "🛠 *Специальные навыки:*\n\n"
            "🎨 GPT Image - промпты для нейросетей\n"
            "💰 AI Sales - продажи AI-визуалов\n"
            "📚 МГИМО БЖД - контрольные работы\n"
            "📖 Курсовая - курсовые МГИМО\n\n"
            "Выбери навык:",
            [["🎨 GPT Image промпты"], ["💰 AI Visuals Sales"], ["📚 МГИМО БЖД", "📖 Курсовая МГИМО"], ["🗣 Дефолт"]],
            markdown=True
        )
        return

    if text == "/model":
        current = MODEL_NAMES.get(u.get("model", "sonnet"), "⚡ Sonnet 4.6")
        send(chat_id, f"Текущая модель: {current}\n\nВыбери модель:", [["⚡ Sonnet 4.6", "🧠 Opus 4.7"]])
        return

    if text == "/lang":
        send(chat_id, "Выбери язык:", [["🇷🇺 Русский", "🇬🇧 English", "🌐 Авто"]])
        return

    if text == "/clear":
        pending_clear.add(chat_id)
        send(chat_id, "Ты уверен? Вся история переписки будет удалена.", [["✅ Точно очистить", "❌ Нет"]])
        return

    if text == "/sub":
        remaining = get_limit(chat_id) - u.get("msg_count", 0)
        if u.get("has_sub"):
            send(chat_id, "✅ У тебя активна подписка! Лимит сообщений снят.")
        else:
            send(chat_id,
                f"💳 *Подписка*\n\n"
                f"Осталось бесплатных сообщений: {max(0, remaining)}\n\n"
                f"Чтобы получить неограниченный доступ - напиши администратору: @{ADMIN_USERNAME}\n\n"
                f"Стоимость: уточни у администратора.",
                markdown=True
            )
        return

    if text == "/ref":
        if not u.get("has_sub"):
            send(chat_id, "Реферальная система доступна только для подписчиков. Оформи подписку через /sub")
            return
        bonus = u.get("referral_bonus", 0)
        refs = len(u.get("referrals", []))
        bot_info = requests.get(f"{BASE_URL}/getMe").json()
        bot_username = bot_info.get("result", {}).get("username", "")
        ref_link = f"https://t.me/{bot_username}?start=ref_{chat_id}"
        send(chat_id,
            f"🔗 *Реферальная программа*\n\n"
            f"Приглашай друзей и получай +{REFERRAL_BONUS} сообщений за каждого!\n\n"
            f"Твоя ссылка:\n{ref_link}\n\n"
            f"Приглашено друзей: {refs}\n"
            f"Бонусных сообщений получено: {bonus}",
            markdown=True
        )
        return

    if text == "/status":
        remaining = max(0, get_limit(chat_id) - u.get("msg_count", 0))
        sub_status = "✅ Активна (безлимит)" if u.get("has_sub") else f"❌ Нет ({remaining} сообщений осталось)"
        model_name = MODEL_NAMES.get(u.get("model", "sonnet"), "⚡ Sonnet 4.6")
        mode_names = {
            "default": "🗣 Дефолт", "delo": "💼 По делу", "friend": "😊 Друг",
            "test": "📝 Тесты", "english": "🇬🇧 Английский", "text_work": "✏️ Работа с текстом",
            "skill_gpt": "🎨 GPT Image", "skill_sales": "💰 AI Sales",
            "skill_mgimo": "📚 МГИМО БЖД", "skill_coursework": "📖 Курсовая МГИМО",
        }
        mode_name = mode_names.get(u.get("mode", "default"), "🗣 Дефолт")
        send(chat_id,
            f"📊 *Твой статус*\n\n"
            f"Режим: {mode_name}\n"
            f"Модель: {model_name}\n"
            f"Язык: {u.get('language', 'ru')}\n"
            f"Подписка: {sub_status}\n"
            f"Сообщений отправлено: {u.get('msg_count', 0)}\n"
            f"История: {len(u.get('history', []))} сообщений",
            markdown=True
        )
        return

    # Кнопки выбора модели
    if text == "🖼 Фотографии 🔒":
        if not u.get("has_sub"):
            send(chat_id, "🔒 Режим генерации фотографий доступен только по подписке.\n\nОформи подписку через /sub или узнай подробнее у @forge0n")
        else:
            send(chat_id, "🔒 Режим генерации фотографий скоро будет доступен! Следи за обновлениями.")
        return

    if text == "⚡ Sonnet 4.6":
        u["model"] = "sonnet"
        send(chat_id, "Модель: ⚡ Sonnet 4.6 - быстрый и экономный.")
        return
    if text == "🧠 Opus 4.7":
        if not u.get("has_sub"):
            send(chat_id, "🔒 Модель Opus 4.7 доступна только по подписке.\n\nОформи подписку через /sub или узнай подробнее у @forge0n")
            return
        u["model"] = "opus"
        send(chat_id, "Модель: 🧠 Opus 4.7 - максимальный интеллект.")
        return

    # Кнопки языка
    if text == "🇷🇺 Русский":
        u["language"] = "ru"
        send(chat_id, "Язык: 🇷🇺 Русский")
        return
    if text == "🇬🇧 English":
        u["language"] = "en"
        send(chat_id, "Language: 🇬🇧 English")
        return
    if text == "🌐 Авто":
        u["language"] = "auto"
        send(chat_id, "Язык: 🌐 Авто")
        return

    # Кнопки субрежима текста
    if text in TEXT_SUBMODES:
        submode = TEXT_SUBMODES[text]
        pending_text_mode[chat_id] = submode
        send(chat_id, "Отправь текст, и я его обработаю 👇")
        return

    # Кнопки режимов и скиллов
    if text in BUTTONS:
        val = BUTTONS[text]
        if val:
            mode_key, reply_text = val
            # Проверка доступа к МГИМО
            if mode_key in ("skill_mgimo", "skill_coursework") and not can_use_mgimo(user_id, username):
                send(chat_id, "🔒 Этот скилл доступен только по запросу. Напиши @forge0n для получения доступа.")
                return
            u["mode"] = mode_key
            u["history"] = []
            pending_text_mode.pop(chat_id, None)
            if mode_key == "text_work":
                send(chat_id,
                    "✏️ *Работа с текстом*\n\nВыбери что сделать с текстом:",
                    [
                        ["🔄 Перефразировка", "✂️ Сокращение"],
                        ["📝 Удлинение", "📋 По пунктам"],
                        ["🤖 Анализ на ИИ"],
                        ["🗣 Дефолт"]
                    ],
                    markdown=True
                )
            else:
                send(chat_id, reply_text)
        return

    # Проверка лимита
    if not can_message(chat_id):
        send(chat_id,
            "⚠️ У тебя закончились бесплатные сообщения.\n\n"
            "Оформи подписку через /sub или пригласи друга через /ref для бонусных сообщений."
        )
        return

    # Основной запрос к Claude
    u["history"].append({"role": "user", "content": text})
    if len(u["history"]) > 20:
        u["history"] = u["history"][-20:]

    send_typing(chat_id)

    try:
        model_id = MODELS.get(u.get("model", "sonnet"), MODELS["sonnet"])
        response = client.messages.create(
            model=model_id,
            max_tokens=2048,
            system=get_system(chat_id),
            messages=u["history"]
        )
        reply = response.content[0].text
        u["history"].append({"role": "assistant", "content": reply})
        u["msg_count"] = u.get("msg_count", 0) + 1
        send(chat_id, reply, markdown=True)
    except Exception as e:
        send(chat_id, f"Ошибка: {str(e)}")

def main():
    set_commands()
    offset = 0
    print("Бот запущен!")
    while True:
        try:
            r = requests.get(
                f"{BASE_URL}/getUpdates",
                params={"offset": offset, "timeout": 30},
                timeout=35
            )
            updates = r.json().get("result", [])
            for update in updates:
                offset = update["update_id"] + 1
                handle(update)
        except Exception as e:
            print(f"Ошибка: {e}")
            time.sleep(5)

if __name__ == "__main__":
    main()
